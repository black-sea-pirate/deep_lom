"""
Document Processing Service

Extracts text from various document formats:
- PDF
- DOCX/DOC
- TXT
- Images (OCR)
"""

import os
from typing import List, Optional
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
import io


class DocumentProcessor:
    """
    Document processor for extracting text from various file formats.
    Supports PDF, DOCX, TXT, and images with OCR.
    """
    
    # Chunk size for splitting text (in characters)
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200
    
    def process_file(self, file_path: str) -> List[str]:
        """
        Process a file and return text chunks.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            List of text chunks
        """
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            text = self._extract_docx(file_path)
        elif file_ext == ".txt":
            text = self._extract_txt(file_path)
        elif file_ext in [".png", ".jpg", ".jpeg"]:
            text = self._extract_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Split into chunks
        return self._split_into_chunks(text)
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n\n".join(paragraphs)
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            return file.read()
    
    def _extract_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    def _split_into_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Uses a simple character-based splitting with overlap
        to maintain context across chunks.
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end of the chunk
            end = start + self.CHUNK_SIZE
            
            # If not at the end, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings
                for sep in [". ", ".\n", "! ", "? ", "\n\n"]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start:
                        end = last_sep + len(sep)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - self.CHUNK_OVERLAP
            if start < 0:
                start = 0
            if start >= len(text):
                break
        
        return chunks


# Global processor instance
document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance."""
    global document_processor
    if document_processor is None:
        document_processor = DocumentProcessor()
    return document_processor
